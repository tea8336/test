# coding:utf-8
# word.py
# yang.wenbo


import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Word:
    'Word类'

    def __init__(self):
        '【初始化】'
        self.office_docx = Document()
        self.str_path = os.path.dirname(os.path.abspath(__file__))

    def create_word(self):
        '【建立Word】'
        # 添加主标题
        self.word_title(input('请输入主标题'))
        # 添加副标题
        self.word_title2(input('请输入副标题'))
        # 添加正文
        while True:
            # 默认段首缩进
            self.word_text('    '+input('请输入正文'))
            # 是否继续添加正文
            if input('按q键保存，按其他任意键继续添加文本').lower() == 'q':
                break
            else:
                pass
        # 保存
        self.word_save('%s\\files\\%s.docx' % (self.str_path, input('请输入保存的文件名')))
        print('保存成功')

    def word_title(self, str_title: str):
        '【主标题】'
        self.office_docx.add_heading(str_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def word_title2(self, str_title2: str):
        '【副标题】'
        self.office_docx.add_paragraph(str_title2, 'Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

    def word_text(self, str_text: str):
        '【正文】'
        self.office_docx.add_paragraph(str_text).alignment = WD_ALIGN_PARAGRAPH.LEFT

    def word_save(self, str_name: str):
        '【保存】'
        self.office_docx.save(str_name)


def main():
    word_obj = Word()
    word_obj.create_word()

if __name__ == '__main__':
    main()
